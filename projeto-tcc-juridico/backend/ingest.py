"""Ingestão de documentos jurídicos: TXT, PDF, DOCX → chunks."""
from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend.config import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    DATA_DIR,
    EXEMPLOS_JSON,
    KNOWLEDGE_DIR,
    TIPOS_DOCUMENTO,
    UPLOADS_DIR,
)


def _classificar_tipo_documento(texto: str, filename: str = "") -> str:
    """
    Classifica o tipo de documento jurídico com base em palavras-chave.
    Retorna: 'peticao_inicial', 'sentenca', 'acordao', 'decisao' ou 'outros'.
    """
    texto_lower = texto.lower()[:3000]  # primeiros 3000 chars são suficientes
    filename_lower = filename.lower()

    # Pontuação por tipo
    scores: dict[str, int] = {}
    for tipo, keywords in TIPOS_DOCUMENTO.items():
        if tipo == "outros":
            continue
        score = 0
        for kw in keywords:
            # Busca no texto e no nome do arquivo
            score += texto_lower.count(kw) * 2
            if kw in filename_lower:
                score += 5
        scores[tipo] = score

    if not scores:
        return "outros"

    melhor_tipo = max(scores, key=scores.get)  # type: ignore[arg-type]
    return melhor_tipo if scores[melhor_tipo] > 0 else "outros"


def _chunk_size_adaptativo(texto: str) -> int:
    """
    Ajusta o chunk size com base no comprimento do texto.
    Textos longos (> 30k chars) usam chunks maiores.
    """
    if len(texto) > 50000:
        return min(CHUNK_SIZE + 600, 1800)
    elif len(texto) > 20000:
        return min(CHUNK_SIZE + 300, 1500)
    return CHUNK_SIZE


def extract_bytes(filename: str | None, content: bytes) -> str:
    if not content:
        raise ValueError("Ficheiro vazio.")
    name = (filename or "").strip().lower()
    if name.endswith(".docx"):
        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(content))
        return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content), strict=False)
        parts = []
        for i, page in enumerate(reader.pages, 1):
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(f"--- Página {i} ---\n{t}")
        return "\n\n".join(parts).strip()
    if name.endswith((".txt", ".md", ".markdown")) or "." not in name:
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return content.decode(enc).strip()
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace").strip()
    raise ValueError(f"Extensão não suportada: {filename!r}")


def extract_path(path: Path) -> str:
    return extract_bytes(path.name, path.read_bytes())


def text_to_documents(
    text: str,
    *,
    source: str,
    extra_meta: dict | None = None,
) -> list[Document]:
    """Texto completo → lista de Document (chunks) com metadata por pedaço."""
    chunk_size = _chunk_size_adaptativo(text)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    # Classifica o tipo de documento
    tipo_doc = _classificar_tipo_documento(text, source)
    meta = {
        "source": source,
        "tipo_documento": tipo_doc,
        **(extra_meta or {}),
    }
    base = Document(page_content=text, metadata=dict(meta))
    chunks = splitter.split_documents([base])
    for i, doc in enumerate(chunks):
        doc.metadata["chunk_index"] = i
        doc.metadata["total_chunks"] = len(chunks)
        doc.metadata["chunk_size_usado"] = chunk_size
    return chunks


def load_exemplo_json() -> list[Document]:
    if not EXEMPLOS_JSON.is_file():
        return []
    raw = json.loads(EXEMPLOS_JSON.read_text(encoding="utf-8"))
    docs: list[Document] = []
    for item in raw.get("peticoes", []):
        texto = (item.get("texto") or "").strip()
        if len(texto) < 80:
            continue
        meta = {
            "source": f"exemplo:{item.get('id', 'sem_id')}",
            "resultado": item.get("resultado", ""),
            "classe": item.get("classe_processual", ""),
            "tipo": "exemplo",
            "tipo_documento": "peticao_inicial",
        }
        docs.extend(text_to_documents(texto, source=meta["source"], extra_meta=meta))
    return docs


def iter_knowledge_files() -> list[Path]:
    paths: list[Path] = []
    for folder in (KNOWLEDGE_DIR, UPLOADS_DIR):
        if not folder.is_dir():
            continue
        for ext in ("*.txt", "*.md", "*.pdf", "*.docx"):
            paths.extend(folder.glob(ext))
    return sorted(set(paths))


def documents_from_path(path: Path) -> list[Document]:
    text = extract_path(path)
    rel = path.relative_to(DATA_DIR) if path.is_relative_to(DATA_DIR) else path.name
    return text_to_documents(text, source=str(rel))


def save_upload(filename: str, content: bytes) -> Path:
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    safe = Path(filename).name or "upload.txt"
    dest = UPLOADS_DIR / safe
    dest.write_bytes(content)
    return dest
